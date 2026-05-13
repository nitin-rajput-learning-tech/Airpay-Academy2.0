#!/usr/bin/env python3
"""Generate the Airpay Academy master documentation .docx from the
concatenated master .md file. Uses python-docx (already installed).

This is intentionally a simple converter — not a full Markdown engine.
It handles: H1/H2/H3 headings, paragraphs, bullet lists, tables
(pipe-delimited), and inline `code` / **bold** / _italic_.
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Airpay brand colors
AIRPAY_PRIMARY = RGBColor(0x00, 0x66, 0xA7)
AIRPAY_ACCENT  = RGBColor(0x0F, 0x7A, 0x73)
TEXT_PRIMARY   = RGBColor(0x1A, 0x1A, 0x2E)
TEXT_SECONDARY = RGBColor(0x5A, 0x60, 0x70)


def add_page_number(paragraph):
    """Insert a PAGE field into a footer paragraph."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def set_doc_defaults(doc):
    """Set Montserrat font and base styling."""
    style = doc.styles['Normal']
    style.font.name = 'Montserrat'
    style.font.size = Pt(10.5)
    style.font.color.rgb = TEXT_PRIMARY
    # East Asian font hint too
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), 'Montserrat')
    rFonts.set(qn('w:hAnsi'), 'Montserrat')

    for h in ('Heading 1', 'Heading 2', 'Heading 3'):
        s = doc.styles[h]
        s.font.name = 'Montserrat'
        s.font.color.rgb = AIRPAY_PRIMARY
        if h == 'Heading 1':
            s.font.size = Pt(20)
            s.font.bold = True
        elif h == 'Heading 2':
            s.font.size = Pt(15)
            s.font.bold = True
        elif h == 'Heading 3':
            s.font.size = Pt(12)
            s.font.bold = True


def parse_inline(text):
    """Return a list of (kind, content) for runs.
    Kinds: 'plain', 'bold', 'italic', 'code'."""
    parts = []
    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`|_[^_]+_)')
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            parts.append(('plain', text[last:m.start()]))
        token = m.group(1)
        if token.startswith('**'):
            parts.append(('bold', token[2:-2]))
        elif token.startswith('`'):
            parts.append(('code', token[1:-1]))
        elif token.startswith('_'):
            parts.append(('italic', token[1:-1]))
        last = m.end()
    if last < len(text):
        parts.append(('plain', text[last:]))
    return parts


def add_inline(paragraph, text):
    """Add inline-formatted runs to a paragraph."""
    for kind, content in parse_inline(text):
        run = paragraph.add_run(content)
        if kind == 'bold':
            run.bold = True
        elif kind == 'italic':
            run.italic = True
        elif kind == 'code':
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = AIRPAY_ACCENT


def add_table(doc, rows):
    """Add a pipe-delimited table. First row treated as header."""
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=n_cols)
    t.style = 'Light Grid Accent 1'
    for i, row in enumerate(rows):
        for j in range(n_cols):
            cell = t.rows[i].cells[j]
            cell.text = ''  # clear
            p = cell.paragraphs[0]
            add_inline(p, row[j] if j < len(row) else '')
            if i == 0:
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = AIRPAY_PRIMARY


def convert(md_path, docx_path):
    text = Path(md_path).read_text(encoding='utf-8')
    lines = text.split('\n')

    doc = Document()
    set_doc_defaults(doc)

    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        # Page numbers in footer
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_number(fp)

    i = 0
    in_code = False
    code_buf = []
    table_buf = []

    def flush_table():
        nonlocal table_buf
        if table_buf:
            # Strip alignment row (---)
            rows = [r for r in table_buf
                    if not all(set(c.strip()) <= set('-: ') for c in r)]
            add_table(doc, rows)
            table_buf = []

    while i < len(lines):
        line = lines[i]

        # Code fence?
        if line.startswith('```'):
            if in_code:
                # close fence
                in_code = False
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_buf))
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = AIRPAY_ACCENT
                code_buf = []
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # Table?
        if '|' in line and line.strip().startswith('|'):
            row = [c.strip() for c in line.strip().strip('|').split('|')]
            table_buf.append(row)
            i += 1
            continue
        flush_table()

        # Headings
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=3)
        elif line.strip() == '---':
            doc.add_paragraph()  # spacing
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            add_inline(p, line[2:].strip())
        elif re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph(style='List Number')
            add_inline(p, re.sub(r'^\d+\.\s', '', line))
        elif line.strip() == '':
            doc.add_paragraph()
        else:
            p = doc.add_paragraph()
            add_inline(p, line)
        i += 1

    flush_table()

    doc.save(docx_path)


if __name__ == '__main__':
    md = sys.argv[1] if len(sys.argv) > 1 else \
        'docs/AIRPAY-ACADEMY-2.0-MASTER-DOCUMENTATION-2026-05-12.md'
    docx = sys.argv[2] if len(sys.argv) > 2 else \
        'docs/AIRPAY-ACADEMY-2.0-MASTER-DOCUMENTATION-2026-05-12.docx'
    convert(md, docx)
    print(f'Generated {docx} from {md}')
